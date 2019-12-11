'''
This is Adade's big final project where I have to write a code to create a Random Taco Recipe Book in a new word
processing document with different pages.
I struggled while writing my code because I wanted to get things write and this is part of my work as a "coder". But
I think, there is always someone who will code better than you.  This is not a reason to give up on your code
especially as Brian's student.  sorry for my "bla bla bla"!
Thanks for taking your time to read my introduction
Let's go in the details of my code.
.

many lines of text
to describe
every little tiny detail that happens in this code
i struggled with
this part was easy
'''

'''
Make sure i have all the librarires I need to use
'''
import docx # import this library to create and update a word document
from docx.shared import Inches # this import allows to get the inches for the width and the height of the picture
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_BREAK # source: stack overflow
import requests # to get the url



'''
I need to create a new word processing document that will have all the information about the Random Taco Cook book that
I need to make.
'''
document = docx.Document() # create a new word processing document

document.add_heading("Random Taco Cookbook", 0) # add the heading Random Taco Cookbook

document.add_picture('Modified_Tai_Taco.jpg', width=Inches(6), height=Inches(4)) # add the picture of the
# Tai Taco in 6 inches with and 4 inches height in order to fit correctly in the page of my world processing document

document.add_paragraph() # add and empty paragraph after the picture to space out the picture and the next paragraph

document.add_paragraph('Credits', 'Heading 3') # add Credits heading to provide the author of the image, the Url source
# and the person who wrote the code to create the document (Adade Gbadoe)

document.add_paragraph('Taco image: Photo by Tai\'s Captures on Unsplash', 'List Bullet') # add the name of
# the image author and add bullet point before ' Taco image: Photo by Tai\'s Captures on Unsplash'

document.add_paragraph('Source: https://taco-1150.herokuapp.com/random/?full_taco=true', 'List Bullet') # add the URL text
# and add bullet point before this paragraph

''''I struggled with the following part. To insert the page break after 'Code by: Adade Gbadoe',I first coded:
document.add_paragraph('Code by: Adade Gbadoe')
 for paragraph in document.paragraphs:
     if 'Adade Gbadoe' in paragraph.text:
        run = paragraph.add_run()
         run.add_break(WD_BREAK.PAGE)
I changed it to the following code later to insert page break after ('Code by: Adade Gbadoe') by using 
paragraph = document.add_paragraph('Code by: Adade Gbadoe') as a reference  before the page break
I have to get rid of: document.add_paragraph('Code by: Adade Gbadoe') in order to avoid the repetition of 
'Code by: Adade Gbadoe' before the following code'''

paragraph = document.add_paragraph('Code by: Adade Gbadoe', 'List Bullet') # add the last paragraph and use it as reference
run = paragraph.add_run() # add run to be able to add the page break
run.add_break(WD_BREAK.PAGE) # add the page break

# working with Brian

url = 'https://taco-1150.herokuapp.com/random/?full_taco=true' # this will give a random recipe for the Tai Taco
random_taco_title = ['First', 'Second', 'Third'] # define this variable for first, second and third taco recipe
# this for loop to pull out each random taco recipe
for i in range(3): # for any random taco recipe in the range(3). range (3) = [0,1,2]
    document.add_heading(f'{random_taco_title[i]} Taco Recipe', 0) # add the heading on each random taco recipe:
    # First Taco Recipe
    # Second Taco Recipe
    # Third Taco Recipe

    response = requests.get(url).json() # provides a dictionary containing a taco recipe. For the next program

    '''many lines of text to describe the recipe from each of the five components of the recipe: seasoning, 
    condiment, mixin, base_layer and shell'''

    seasoning = response['seasoning']['recipe'] # get the seasoning's recipe
    seasoning_name = response['seasoning']['name'] # get the seasoning name

    condiment = response['condiment']['recipe'] # get the condiment's recipe
    condiment_name = response['condiment']['name'] # get the condiment'name

    mixin = response['mixin']['recipe'] # get the mixin's recipe
    mixin_name = response['mixin']['name'] # get the mixin's name

    base_layer = response['base_layer']['recipe'] # get the base-layer's recipe
    base_layer_name = response['base_layer']['name'] # get the base-layer's name

    shell = response['shell']['recipe'] # get the shell's recipe
    shell_name = response['shell']['name'] # get the shell's name

#Add the name of each component of the taco  with 'Heading' style 3

    document.add_paragraph(seasoning_name, 'Heading 3') # add the name of the seasoning in heading style 3
    document.add_paragraph(seasoning) # add the seasoning paragraph

    document.add_paragraph(condiment_name, 'Heading 3')# add the name of the condiment in heading style 3
    document.add_paragraph(condiment)# add the condiment paragraph

    document.add_paragraph(mixin_name, 'Heading 3')# add the name of the mixin in heading style 3
    document.add_paragraph(mixin) # add the mixin paragraph

    document.add_paragraph(base_layer_name, 'Heading 3') # add the name of the base-layer in heading style 3
    document.add_paragraph(base_layer) # add the base-layer paragraph

    document.add_paragraph(shell_name, 'Heading 3') # add the name of the shell in heading style 3
    paragraph = document.add_paragraph(shell) # add the shell paragraph and
    # use this paragraph as reference to add a page break after each random taco recipe
    run = paragraph.add_run() # add run to be able to add the page break
    run.add_break(WD_BREAK.PAGE) # insert the page break


document.save('First_Page.docx') # save the world document created